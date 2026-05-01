# document/docx_watermark.py
"""
Word文档水印处理器 - 零宽字符隐写

本模块实现 Word 文档的隐形水印：
- 嵌入方式：零宽字符（Zero-Width Characters）
- 优点：完全不可见，不影响文档内容
- 缺点：可能被某些编辑器清理

技术方案：
1. 将水印文本转换为 UTF-8 字节
2. 每个字节转换为 8 位二进制
3. 每位二进制映射为零宽字符：
   - 0 → U+200B (零宽空格 ZWSP)
   - 1 → U+200C (零宽非连接符 ZWNJ)
4. 将零宽字符序列插入文档末尾

编码示例:
    原文: "A" (ASCII 65)
    二进制: 01000001
    零宽字符: ZWSP ZWNJ ZWSP ZWSP ZWSP ZWSP ZWSP ZWNJ

使用示例:
    >>> handler = DOCXHandler()
    >>> 
    >>> # 嵌入水印
    >>> result = handler.embed(
    ...     "document.docx",
    ...     WatermarkData(content="机密文档"),
    ...     "output.docx"
    ... )
    >>> 
    >>> # 提取水印
    >>> result = handler.extract("output.docx")
    >>> print(result.watermark.content)

依赖:
    - python-docx: Word 文档操作

Author: StealthMark Team
Date: 2026-04-28
"""

import os
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any
import logging

# 延迟导入
try:
    import docx
except ImportError:
    docx = None

from ..core.base import (
    BaseHandler, WatermarkData, WatermarkStatus,
    EmbedResult, ExtractResult, VerifyResult
)

# 模块日志
logger = logging.getLogger(__name__)


# ==================== 零宽字符定义 ====================

ZWSP = '\u200b'      # 零宽空格 (Zero-Width Space) - 表示二进制 0
ZWNJ = '\u200c'      # 零宽非连接符 (Zero-Width Non-Joiner) - 表示二进制 1


class DOCXHandler(BaseHandler):
    """
    Word文档水印处理器
    
    使用零宽字符实现隐形水印嵌入。
    
    零宽字符特性:
        - 在屏幕和打印中不可见
        - 不占用视觉空间
        - 可被某些文本编辑器清理
        - 在 Word 中稳定保存
    
    容量计算:
        每个字符需要 8 个零宽字符（1字节 = 8位）
        例：10个中文字符 ≈ 30字节 ≈ 240个零宽字符
    
    Attributes:
        无特殊属性，继承自 BaseHandler
    
    限制:
        - 某些编辑器（如记事本）可能删除零宽字符
        - 文档另存为其他格式时可能丢失
    """
    
    SUPPORTED_EXTENSIONS = ('.docx',)
    HANDLER_NAME = "docx"
    
    # ==================== 零宽字符编解码 ====================
    
    def _text_to_zwc(self, text: str) -> str:
        """
        将文本转换为零宽字符序列
        
        转换流程:
        1. UTF-8 编码文本 → 字节序列
        2. 每个字节 → 8位二进制
        3. 每位 → 零宽字符
        
        Args:
            text: 原始文本
        
        Returns:
            str: 零宽字符序列（长度 = len(text.encode('utf-8')) * 8）
        
        Example:
            >>> zwc = handler._text_to_zwc("A")
            >>> len(zwc)  # 8个零宽字符
            8
        """
        logger.debug(f"Converting text to ZWC: {len(text)} chars")
        
        # UTF-8 编码
        bytes_data = text.encode('utf-8')
        
        # 转换为零宽字符
        zwc_text = []
        for byte in bytes_data:
            # 每个字节转换为8位，高位在前
            for i in range(7, -1, -1):
                bit = (byte >> i) & 1
                zwc_text.append(ZWSP if bit == 0 else ZWNJ)
        
        result = ''.join(zwc_text)
        logger.debug(f"ZWC sequence length: {len(result)}")
        
        return result
    
    def _zwc_to_text(self, zwc_seq: str) -> Optional[str]:
        """
        将零宽字符序列还原为文本
        
        还原流程:
        1. 提取零宽字符（忽略其他字符）
        2. 每8个零宽字符 → 1字节
        3. UTF-8 解码字节序列
        
        Args:
            zwc_seq: 零宽字符序列（可能混有其他字符）
        
        Returns:
            Optional[str]: 还原的文本，失败返回 None
        
        Note:
            序列中的非零宽字符会被忽略，允许从混合文本中提取。
        """
        logger.debug(f"Converting ZWC to text: {len(zwc_seq)} chars")
        
        try:
            # 提取零宽字符对应的位
            bits = []
            for char in zwc_seq:
                if char in (ZWSP, ZWNJ):
                    bits.append('1' if char != ZWSP else '0')
            
            # 每8位转换为1字节
            bytes_data = []
            for i in range(0, len(bits), 8):
                byte_bits = bits[i:i+8]
                if len(byte_bits) < 8:
                    break  # 不足8位，丢弃
                byte = int(''.join(byte_bits), 2)
                bytes_data.append(byte)
            
            # UTF-8 解码
            result = bytes(bytes_data).decode('utf-8')
            logger.debug(f"Decoded text: {len(result)} chars")
            
            return result
            
        except Exception as e:
            logger.error(f"ZWC decode failed: {e}")
            return None
    
    # ==================== 水印操作 ====================
    
    def embed(self, file_path: str, watermark: WatermarkData,
              output_path: str, **kwargs) -> EmbedResult:
        """
        嵌入水印到 Word 文档
        
        嵌入流程:
        1. 打开 Word 文档
        2. 将水印转换为零宽字符序列
        3. 在文档末尾添加一个包含零宽字符的 run
        4. 设置 run 字体大小为最小（1磅）
        5. 保存文档
        
        Args:
            file_path: 原始 Word 文档路径
            watermark: 水印数据对象
            output_path: 输出文档路径
            **kwargs: 额外参数（未使用）
        
        Returns:
            EmbedResult: 嵌入结果
        
        Note:
            零宽字符以最小字号添加，即使显示也不会明显。
        """
        logger.info(f"DOCX embed: {file_path} -> {output_path}")
        
        # 检查依赖
        if docx is None:
            logger.error("python-docx not installed")
            return EmbedResult(
                status=WatermarkStatus.FAILED,
                message="需要安装python-docx库: pip install python-docx",
                file_path=file_path
            )
        
        # 验证文件
        error_result = self._validate_file(file_path)
        if error_result:
            return error_result
        
        try:
            # 打开文档
            doc = docx.Document(file_path)
            
            # 转换为零宽字符
            zwc_seq = self._text_to_zwc(watermark.content)
            
            # 添加到文档末尾（最后一个段落）
            if doc.paragraphs:
                last_para = doc.paragraphs[-1]
                run = last_para.add_run(zwc_seq)
            else:
                # 空文档，添加新段落
                para = doc.add_paragraph()
                run = para.add_run(zwc_seq)
            
            # 设置最小字号（1磅）
            from docx.oxml.ns import qn
            rPr = run._element.makeelement(qn('w:rPr'), {})
            sz = rPr.makeelement(qn('w:sz'), {})
            sz.set(qn('w:val'), '1')  # 1磅 = 2半磅
            rPr.append(sz)
            run._element.append(rPr)
            
            # 保存
            doc.save(output_path)
            
            logger.info(f"DOCX embed success: {output_path}")
            return self._create_success_result(output_path)
            
        except Exception as e:
            logger.error(f"DOCX embed failed: {e}")
            return EmbedResult(
                status=WatermarkStatus.FAILED,
                message=f"嵌入失败: {str(e)}",
                file_path=file_path
            )
    
    def extract(self, file_path: str, **kwargs) -> ExtractResult:
        """
        从 Word 文档提取水印
        
        提取流程:
        1. 解压 docx 文件（ZIP 格式）
        2. 读取 word/document.xml
        3. 提取所有 <w:t> 元素的文本
        4. 从文本中提取零宽字符序列
        5. 解码为零宽字符
        
        Args:
            file_path: Word 文档路径
            **kwargs: 额外参数（未使用）
        
        Returns:
            ExtractResult: 提取结果
        
        Note:
            直接解析 XML 而非使用 python-docx，更可靠。
        """
        logger.info(f"DOCX extract: {file_path}")
        
        # 检查依赖
        if docx is None:
            logger.error("python-docx not installed")
            return ExtractResult(
                status=WatermarkStatus.FAILED,
                message="需要安装python-docx库",
                file_path=file_path
            )
        
        # 验证文件
        error_result = self._validate_file(file_path)
        if error_result:
            return ExtractResult(
                status=error_result.status,
                message=error_result.message,
                file_path=file_path
            )
        
        try:
            # 解压 docx（ZIP 格式）
            with zipfile.ZipFile(file_path, 'r') as zf:
                with zf.open('word/document.xml') as f:
                    content = f.read().decode('utf-8')
            
            # 解析 XML
            root = ET.fromstring(content)
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # 提取所有文本
            all_text = []
            for t_elem in root.iter(f'{{{ns}}}t'):
                if t_elem.text:
                    all_text.append(t_elem.text)
            
            full_text = ''.join(all_text)
            
            # 提取零宽字符
            zwc_chars = []
            for char in full_text:
                if char in (ZWSP, ZWNJ):
                    zwc_chars.append(char)
            
            zwc_seq = ''.join(zwc_chars)
            
            # 检查是否找到水印
            if not zwc_seq:
                logger.warning("No ZWC found in document")
                return ExtractResult(
                    status=WatermarkStatus.EXTRACTION_FAILED,
                    message="未找到水印",
                    file_path=file_path
                )
            
            # 解码
            text = self._zwc_to_text(zwc_seq)
            
            if text:
                logger.info(f"DOCX extract success: {text[:30]}...")
                return ExtractResult(
                    status=WatermarkStatus.SUCCESS,
                    message="水印提取成功",
                    file_path=file_path,
                    watermark=WatermarkData(content=text)
                )
            else:
                logger.warning("ZWC decode failed")
                return ExtractResult(
                    status=WatermarkStatus.EXTRACTION_FAILED,
                    message="水印解码失败",
                    file_path=file_path
                )
                
        except Exception as e:
            logger.error(f"DOCX extract failed: {e}")
            return ExtractResult(
                status=WatermarkStatus.EXTRACTION_FAILED,
                message=f"提取失败: {str(e)}",
                file_path=file_path
            )
    
    def verify(self, file_path: str, original_watermark: WatermarkData,
               **kwargs) -> VerifyResult:
        """
        验证 Word 文档水印
        
        Args:
            file_path: Word 文档路径
            original_watermark: 原始水印数据
            **kwargs: 额外参数
        
        Returns:
            VerifyResult: 验证结果
        """
        logger.info(f"DOCX verify: {file_path}")
        
        # 提取水印
        extract_result = self.extract(file_path)
        
        # 提取失败
        if not extract_result.is_success or not extract_result.watermark:
            return VerifyResult(
                status=extract_result.status,
                is_valid=False,
                is_integrity_ok=False,
                match_score=0.0,
                message=f"提取失败: {extract_result.message}"
            )
        
        # 比对
        extracted = extract_result.watermark.content
        original = original_watermark.content
        
        is_match = extracted == original
        match_score = 1.0 if is_match else 0.0
        
        result = VerifyResult(
            status=WatermarkStatus.SUCCESS if is_match else WatermarkStatus.VERIFICATION_FAILED,
            is_valid=is_match,
            is_integrity_ok=True,
            match_score=match_score,
            message="验证通过" if is_match else "水印不匹配"
        )
        
        logger.info(f"DOCX verify result: valid={is_match}")
        return result


# 模块初始化日志
logger.info(f"{__name__} module loaded - DOCX handler ready")
